"""
Generates resume.docx — Word version of Rumana Appi's resume.
Run: pip install python-docx && python generate_word_resume.py
Output: resume.docx in the same folder
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE  = os.path.dirname(os.path.abspath(__file__))
OUT   = os.path.join(BASE, "resume.docx")

# ── COLOR PALETTE ─────────────────────────────────────────────────────────────
ACCENT   = RGBColor(0, 79, 144)      # #004F90  — section headers, title
BLACK    = RGBColor(17, 17, 17)      # #111111  — body text
GRAY     = RGBColor(80, 80, 80)      # #505050  — company / dates
WHITE    = RGBColor(255, 255, 255)

# ── HELPERS ───────────────────────────────────────────────────────────────────

def set_font(run, size_pt, bold=False, color=None, italic=False):
    run.bold    = bold
    run.italic  = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color

def add_para(doc, text='', alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        p.add_run(text)
    return p

def add_hr(doc, color_hex='004F90'):
    """Add a horizontal rule (colored bottom border on an empty paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    run.font.name = 'Calibri'
    add_hr(doc)

def add_job(doc, title, company, dates, employment_type=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)

    # Title (bold)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = BLACK

    # Dates (right-aligned via tab)
    tab_stop_pos = Inches(6.3)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(tab_stop_pos.pt * 20)))
    tabs.append(tab)
    pPr.append(tabs)
    r_tab = p.add_run('\t')
    r_tab.font.size = Pt(10.5)
    r_dates = p.add_run(dates)
    r_dates.font.size = Pt(10)
    r_dates.font.color.rgb = GRAY

    # Company line
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    r_co = p2.add_run(company)
    r_co.italic = True
    r_co.font.size = Pt(10)
    r_co.font.color.rgb = GRAY
    if employment_type:
        r_et = p2.add_run(f'  ·  {employment_type}')
        r_et.italic = True
        r_et.font.size = Pt(9.5)
        r_et.font.color.rgb = GRAY

def add_bullet(doc, text_parts, indent_level=1):
    """text_parts: list of (text, bold) tuples"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 * indent_level)
    for text, bold in text_parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK
        run.font.name = 'Calibri'
    return p

def add_project(doc, title, tech, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = BLACK
    r2 = p.add_run(f'  |  {tech}')
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY
    for b in bullets:
        add_bullet(doc, b)

def bold_phrases(text, phrases):
    """Split text into (segment, bold) parts for bold phrase highlighting."""
    parts = []
    remaining = text
    while remaining:
        earliest = None
        match_phrase = None
        for phrase in phrases:
            idx = remaining.find(phrase)
            if idx != -1 and (earliest is None or idx < earliest):
                earliest = idx
                match_phrase = phrase
        if earliest is None:
            parts.append((remaining, False))
            break
        if earliest > 0:
            parts.append((remaining[:earliest], False))
        parts.append((match_phrase, True))
        remaining = remaining[earliest + len(match_phrase):]
    return parts

# ── BUILD DOCUMENT ────────────────────────────────────────────────────────────

doc = Document()

# ── Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# ── Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── List Bullet style tweak
lb = doc.styles['List Bullet']
lb.font.name = 'Calibri'
lb.font.size = Pt(10)

# ─────────────────────────────────────────────────────────────────────────────
# HEADING
# ─────────────────────────────────────────────────────────────────────────────

# Name
p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_name.paragraph_format.space_before = Pt(0)
p_name.paragraph_format.space_after  = Pt(2)
r = p_name.add_run('Rumana Appi')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(17, 17, 17)
r.font.name = 'Calibri'

# Job title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(4)
r = p_title.add_run('Senior Data & Business Analyst')
r.font.size = Pt(13)
r.font.color.rgb = ACCENT
r.font.name = 'Calibri'

# Contact line
p_contact = doc.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contact.paragraph_format.space_before = Pt(0)
p_contact.paragraph_format.space_after  = Pt(6)
contacts = [
    ('📞 +91 7766829963', False),
    ('   |   ', False),
    ('✉ rumanappi125@gmail.com', False),
    ('   |   ', False),
    ('🔗 linkedin.com/in/rumana-appi', False),
    ('   |   ', False),
    ('🌐 rumanappi125.github.io/rumana-portfolio', False),
]
for text, bold in contacts:
    r = p_contact.add_run(text)
    r.bold = bold
    r.font.size = Pt(9.5)
    r.font.color.rgb = GRAY
    r.font.name = 'Calibri'

# ─────────────────────────────────────────────────────────────────────────────
# PROFILE SUMMARY
# ─────────────────────────────────────────────────────────────────────────────

add_section_heading(doc, 'Profile Summary')
p_sum = doc.add_paragraph()
p_sum.paragraph_format.space_before = Pt(3)
p_sum.paragraph_format.space_after  = Pt(3)
summary_parts = bold_phrases(
    "Senior Data & Business Analyst with 3+ years delivering analytics at scale — from "
    "multi-terabyte ETL pipelines (PySpark + AWS Redshift) to executive Power BI dashboards "
    "with 100+ KPIs. Domain expertise across e-commerce (Target), insurance (Max Life Insurance) "
    "and EdTech (Physics Wallah). Known for turning raw data into decisions: 15% checkout lift, "
    "12% policy renewal improvement, and 40%+ pipeline efficiency gains across 10M+ row datasets.",
    ["3+ years", "PySpark + AWS Redshift", "Power BI dashboards", "100+ KPIs",
     "15% checkout lift", "12% policy renewal improvement", "40%+ pipeline efficiency gains"]
)
for text, bold in summary_parts:
    r = p_sum.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK
    r.font.name = 'Calibri'

# ─────────────────────────────────────────────────────────────────────────────
# TECHNICAL SKILLS
# ─────────────────────────────────────────────────────────────────────────────

add_section_heading(doc, 'Technical Skills')

skills = [
    ("Query Languages & Databases",
     "SQL (Advanced), MySQL, SQL Server, PostgreSQL, CTEs, Window Functions "
     "(ROW_NUMBER, RANK, LAG, LEAD), Joins, Indexing, Stored Procedures, Views, "
     "Query Optimization, AWS Redshift"),
    ("Data Engineering & Big Data",
     "PySpark, Apache Spark, ETL/ELT Pipeline Design, Data Partitioning, Query Tuning, "
     "Multi-Terabyte Dataset Processing, JSON/CSV Ingestion, Data Quality Frameworks "
     "(Null Checks, Range Validation, Deduplication)"),
    ("Python & Analytics",
     "Pandas, NumPy, Matplotlib, Seaborn, EDA, Feature Engineering, "
     "Outlier Detection, Statistical Analysis, Probability, Predictive Analytics"),
    ("Business Intelligence & Visualization",
     "Power BI (DAX, Power Query, Row Level Security, Data Modelling), Metabase, "
     "Qlik Sense, Excel (Advanced), Google Sheets, KPI Dashboard Design, "
     "Product Funnel Analysis, Cohort Analysis"),
    ("Data Modelling & Architecture",
     "Star Schema, Snowflake Schema, Dimensional Modelling, "
     "AI/ML Training Dataset Design, Real-Time Analytics Architecture"),
    ("Cloud & Platforms",
     "AWS (Redshift, S3), Databricks, Git, JIRA, Confluence, Agile/Scrum Methodology"),
    ("Business & Soft Skills",
     "Stakeholder Management, Requirements Gathering, BRD/FRD Documentation, "
     "Cross-functional Collaboration, Executive Reporting, KPI Definition & Tracking"),
]

for category, items in skills:
    p_sk = doc.add_paragraph()
    p_sk.paragraph_format.space_before = Pt(2)
    p_sk.paragraph_format.space_after  = Pt(2)
    r1 = p_sk.add_run(f'{category}:  ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = BLACK
    r2 = p_sk.add_run(items)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(50, 50, 50)

# ─────────────────────────────────────────────────────────────────────────────
# WORK EXPERIENCE
# ─────────────────────────────────────────────────────────────────────────────

add_section_heading(doc, 'Work Experience')

# ── Job 1
add_job(doc, 'Senior Data Analyst', 'E-Mech Solutions, Gurugram (Client: Target)',
        'Jan 2025 – Present', 'Full-Time')
exp1 = [
    ([("Designed and optimized complex SQL queries (CTEs, window functions, indexing strategies), reducing query execution time by ", False), ("35%", True), (" across large datasets (10M+ records).", False)]),
    ([("Architected high-performance data pipelines using partitioning and query tuning, improving data retrieval efficiency by ", False), ("40%+", True), (" across multi-terabyte datasets.", False)]),
    ([("Designed and implemented dimensional data models (Star/Snowflake schemas) to support AI/ML training datasets and real-time analytics.", False)]),
    ([("Built a Product Funnel Dashboard tracking user journeys across signup → checkout → payment, identifying drop-off stages and increasing checkout completions by ", False), ("15%.", True)]),
    ([("Identified and defined ", False), ("100+ KPIs", True), (" (DAU, MAU, Revenue, Retention, Churn, Transactions, MoM/YoY/WoW changes) across multiple business categories.", False)]),
    ([("Engineered scalable data preprocessing pipelines in Python (Pandas, NumPy) for structured and semi-structured datasets (JSON, CSV).", False)]),
    ([("Automated reporting processes, recovering ", False), ("9 hours per week", True), (" in operational overhead.", False)]),
]
for b in exp1: add_bullet(doc, b)

# ── Job 2
add_job(doc, 'MIS & Business Analyst', 'Randstad, Gurugram (Client: Max Life Insurance)',
        'Jul 2024 – Dec 2024', 'Full-Time')
exp2 = [
    ([("Collaborated with cross-functional teams to translate business requirements into analytics strategies, improving project efficiency by ", False), ("15%.", True)]),
    ([("Leveraged SQL, Power BI, and Excel to extract insights from 10M+ row datasets, driving a ", False), ("30%", True), (" improvement in decision-making speed.", False)]),
    ([("Compiled complex SQL queries using Joins, Window Functions, CTEs, Views, and Stored Procedures, enhancing data processing efficiency by ", False), ("10%.", True)]),
    ([("Designed automated MIS reports for policy renewals and agent performance tracking, reducing manual effort by ", False), ("40%", True), (" and enabling real-time senior management visibility.", False)]),
    ([("Conducted customer churn and retention analysis using cohort modelling, contributing to a ", False), ("12%", True), (" improvement in policy renewal rates.", False)]),
]
for b in exp2: add_bullet(doc, b)

# ── Job 3
add_job(doc, 'Data Analyst (Associate)', 'PW (Physics Wallah), Noida', 'Apr 2023 – Nov 2023')
exp3 = [
    ([("Cleaned and preprocessed data from multiple sources (Excel, CSV, SQL Server), reducing processing time by ", False), ("20%", True), (" and increasing efficiency by ", False), ("20%.", True)]),
    ([("Developed ", False), ("20+ key business metrics", True), (" and interactive dashboards using Power BI and Qlik Sense, supporting quarterly KPI achievement.", False)]),
    ([("Analysed ", False), ("50K+ learner records", True), (" to identify drop-off patterns, enabling content redesign that improved course completion rates by ", False), ("18%.", True)]),
    ([("Built automated data reconciliation workflows between CRM, LMS, and billing systems, eliminating discrepancies and saving ", False), ("6+ hours", True), (" of manual effort weekly.", False)]),
]
for b in exp3: add_bullet(doc, b)

# ── Job 4
add_job(doc, 'Data Analyst Intern', 'I-Neuron, Bengaluru', 'Mar 2022 – Mar 2023')
exp4 = [
    ([("Implemented data quality rules (null checks, range validation, deduplication), improving AI training data trust by ", False), ("25%.", True)]),
    ([("Automated data validation and reconciliation frameworks, reducing data inconsistencies by ", False), ("30%", True), (" in production environments.", False)]),
]
for b in exp4: add_bullet(doc, b)

# ─────────────────────────────────────────────────────────────────────────────
# PROJECTS
# ─────────────────────────────────────────────────────────────────────────────

add_section_heading(doc, 'Projects')

add_project(doc,
    'NREGA Employment Analytics Dashboard',
    'Power BI · DAX · Python · Star Schema',
    [
        [("Built a ", False), ("5-page Power BI dashboard", True), (" analysing 15 states and 75 districts across FY2020–FY2024 — covering employment demand, social equity (SC/ST/Women), wage compliance, and works completion.", False)],
        [("Designed star-schema model with ", False), ("34 DAX measures", True), (" including person-days, demand fulfilment rate, 15-day wage payment %, YoY growth. Implemented Row Level Security by state and district.", False)],
        [("Portfolio: rumanappi125.github.io/rumana-portfolio  |  GitHub: github.com/rumanappi125/rumana-portfolio", False)],
    ]
)

add_project(doc,
    'EdTech Student Engagement Dashboard',
    'Power BI · DAX · Python · PW Branded Theme',
    [
        [("Built a ", False), ("5-page Power BI dashboard", True), (" for Physics Wallah analysing 50K+ learner records across 15 courses — tracking engagement, completion, drop-off, assessment scores, and content effectiveness.", False)],
        [("Developed ", False), ("35 DAX measures", True), (" (completion rate, active rate, drop-off funnel, engagement score) with 5 dropdown slicers per page for Grade, Subject, City Tier, Subscription, and Academic Year.", False)],
        [("Portfolio: rumanappi125.github.io/rumana-portfolio  |  GitHub: github.com/rumanappi125/rumana-portfolio", False)],
    ]
)

add_project(doc,
    'COVID-19 Epidemic Analysis',
    'MySQL · CTEs · Window Functions · Stored Procedures',
    [
        [("Processed and analysed ", False), ("1.69B+ records", True), (" using advanced MySQL — 7-day rolling averages, WoW growth rates, CFR categorisation, continent summaries, and country recovery rankings using CTEs and window functions.", False)],
        [("GitHub: github.com/rumanappi125/rumana-portfolio/tree/master/sql-projects/covid-analysis", False)],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# CERTIFICATIONS
# ─────────────────────────────────────────────────────────────────────────────

add_section_heading(doc, 'Certifications')
p_cert = doc.add_paragraph()
p_cert.paragraph_format.space_before = Pt(4)
p_cert.paragraph_format.space_after  = Pt(2)
r1 = p_cert.add_run('Databricks Certified Data Analyst Associate')
r1.bold = True
r1.font.size = Pt(10.5)
r1.font.color.rgb = BLACK
r2 = p_cert.add_run('  ·  Databricks')
r2.font.size = Pt(10)
r2.font.color.rgb = GRAY

# ─────────────────────────────────────────────────────────────────────────────
# EDUCATION
# ─────────────────────────────────────────────────────────────────────────────

add_section_heading(doc, 'Education')
add_job(doc,
    'Bachelor of Technology – Mechanical Engineering',
    'Netaji Subhash Institute of Technology (NSIT)',
    'Jun 2015 – Aug 2019')
p_edu = doc.add_paragraph()
p_edu.paragraph_format.space_before = Pt(0)
p_edu.paragraph_format.space_after  = Pt(2)
r_edu = p_edu.add_run('CGPA: 8.2  |  First Class with Distinction')
r_edu.italic = True
r_edu.font.size = Pt(10)
r_edu.font.color.rgb = GRAY

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────

doc.save(OUT)
print(f"resume.docx saved to: {OUT}")
